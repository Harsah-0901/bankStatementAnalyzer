import os
import re
import json
import pandas as pd
import PyPDF2
import docx
import google.generativeai as genai
from dotenv import load_dotenv
from pathlib import Path

load_dotenv(dotenv_path=Path(".env"))
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

class BankStatementProcessor:
    def __init__(self):
        self.categories = [
            "Utilities", "Food & Dining", "Travel & Transportation",
            "Subscriptions", "EMIs or Loans", "Shopping",
            "Healthcare", "Miscellaneous"
        ]
        self.model = genai.GenerativeModel("models/gemini-1.5-pro")

    def extract_from_pdf(self, file_path):
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text

    def extract_from_excel(self, file_path):
        df = pd.read_excel(file_path)
        return df.to_string()

    def extract_from_csv(self, file_path):
        df = pd.read_csv(file_path)
        return df.to_string()

    def extract_from_word(self, file_path):
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text += "\t".join(row_text) + "\n"
        return text

    def extract_text_from_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.extract_from_pdf(file_path)
        elif ext in [".xlsx", ".xls"]:
            return self.extract_from_excel(file_path)
        elif ext == ".csv":
            return self.extract_from_csv(file_path)
        elif ext in [".docx", ".doc"]:
            return self.extract_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def parse_transactions(self, text):
        prompt = f"""
        Extract transaction data from the following bank statement text. 
        For each transaction, provide: date, description, amount, and type (credit/debit).
        Format the output as a JSON array.

        Bank statement text:
        {text[:5000]}

        JSON Format:
        [
          {{
            "date": "YYYY-MM-DD",
            "description": "Transaction description",
            "amount": 123.45,
            "type": "credit or debit"
          }}
        ]
        """
        response = self.model.generate_content(prompt)
        result = response.text
        match = re.search(r'\[.*\]', result, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except:
                return []
        return []

    def categorize_transactions(self, transactions):
        categorized = []
        for batch in self._batch_transactions(transactions, 20):
            prompt = f"""
            Categorize the following transactions into these categories:
            {', '.join(self.categories)}

            Transactions:
            {json.dumps(batch, indent=2)}

            For each transaction, add a "category" field.
            Return the result as a JSON array.
            """
            response = self.model.generate_content(prompt)
            result = response.text
            match = re.search(r'\[.*\]', result, re.DOTALL)
            if match:
                try:
                    batch_categorized = json.loads(match.group())
                    categorized.extend(batch_categorized)
                except:
                    pass
        return categorized

    def _batch_transactions(self, transactions, batch_size):
        for i in range(0, len(transactions), batch_size):
            yield transactions[i:i + batch_size]

    def generate_spending_summary(self, categorized_transactions):
        summary = {cat: 0 for cat in self.categories}
        for txn in categorized_transactions:
            cat = txn.get("category", "Miscellaneous")
            if txn["type"] == "debit" and isinstance(txn["amount"], (int, float)):
                summary[cat] += txn["amount"]
        return summary

    def process_file(self, file_path):
        text = self.extract_text_from_file(file_path)
        transactions = self.parse_transactions(text)
        categorized = self.categorize_transactions(transactions)
        summary = self.generate_spending_summary(categorized)
        return {
            "transactions": categorized,
            "summary": summary
        }
