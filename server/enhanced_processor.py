import os
import re
import json
import pandas as pd
import PyPDF2
import docx
import google.generativeai as genai
from dotenv import load_dotenv
from pathlib import Path
from db_connection import execute_query
import datetime

# Load environment variables
load_dotenv(dotenv_path=Path(".env"))
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

class BankStatementProcessor:
    def __init__(self):
        self.model = genai.GenerativeModel("models/gemini-1.5-pro")
        # Get categories from database
        self.categories = self._get_categories_from_db()
        
    def _get_categories_from_db(self):
        """Fetch categories from the database"""
        try:
            categories_data = execute_query(
                "SELECT id, name FROM categories",
                fetch=True
            )
            return {cat['name']: cat['id'] for cat in categories_data}
        except Exception as e:
            print(f"Error fetching categories: {e}")
            # Fallback to default categories
            return {
                "Utilities": 1,
                "Food & Dining": 2,
                "Travel & Transportation": 3,
                "Subscriptions": 4,
                "EMIs or Loans": 5,
                "Shopping": 6,
                "Healthcare": 7,
                "Miscellaneous": 8
            }
            
    def extract_from_pdf(self, file_path):
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text

    def extract_from_excel(self, file_path):
        df = pd.read_excel(file_path)
        return df.to_string()

    def extract_from_csv(self, file_path):
        df = pd.read_csv(file_path)
        return df.to_string()

    def extract_from_word(self, file_path):
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text += "\t".join(row_text) + "\n"
        return text

    def extract_text_from_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.extract_from_pdf(file_path)
        elif ext in [".xlsx", ".xls"]:
            return self.extract_from_excel(file_path)
        elif ext == ".csv":
            return self.extract_from_csv(file_path)
        elif ext in [".docx", ".doc"]:
            return self.extract_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def parse_transactions(self, text):
        prompt = f"""
        Extract transaction data from the following bank statement text. 
        For each transaction, provide: date, description, amount, and type (credit/debit).
        Try to determine the date format used and standardize it to YYYY-MM-DD format.
        If the transaction is money spent/going out, mark it as "debit".
        If the transaction is money received/coming in, mark it as "credit".
        Format the output as a JSON array.
        Bank statement text:
        {text[:5000]}
        JSON Format:
        [
          {{
            "date": "YYYY-MM-DD",
            "description": "Transaction description",
            "amount": 123.45,
            "type": "credit or debit"
          }}
        ]
        """
        response = self.model.generate_content(prompt)
        result = response.text
        match = re.search(r'\[.*\]', result, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except:
                return []
        return []

    def categorize_transactions(self, transactions):
        categorized = []
        category_names = list(self.categories.keys())
        
        for batch in self._batch_transactions(transactions, 20):
            prompt = f"""
            Categorize the following transactions into these categories:
            {', '.join(category_names)}
            Transactions:
            {json.dumps(batch, indent=2)}
            For each transaction, add a "category" field. Choose the most appropriate category based on the transaction description.
            Return the result as a JSON array.
            """
            response = self.model.generate_content(prompt)
            result = response.text
            match = re.search(r'\[.*\]', result, re.DOTALL)
            if match:
                try:
                    batch_categorized = json.loads(match.group())
                    categorized.extend(batch_categorized)
                except:
                    pass
        return categorized

    def _batch_transactions(self, transactions, batch_size):
        for i in range(0, len(transactions), batch_size):
            yield transactions[i:i + batch_size]

    def generate_spending_summary(self, categorized_transactions):
        summary = {cat: 0 for cat in self.categories.keys()}
        for txn in categorized_transactions:
            cat = txn.get("category", "Miscellaneous")
            if txn.get("type", "").lower() == "debit":
                try:
                    amount = float(txn["amount"])
                    summary[cat] += amount
                except (ValueError, TypeError):
                    continue
        return summary

    def process_file(self, file_path, user_id, statement_name=None):
        """Process a bank statement file and save results to database"""
        try:
            # Extract text from file
            text = self.extract_text_from_file(file_path)
            
            # Create statement record
            statement_id = execute_query(
                """INSERT INTO statements 
                   (user_id, statement_name, file_name, processing_status) 
                   VALUES (%s, %s, %s, 'pending')""",
                (user_id, statement_name or os.path.basename(file_path), os.path.basename(file_path))
            )
            
            # Parse transactions
            transactions = self.parse_transactions(text)
            
            # Categorize transactions
            categorized = self.categorize_transactions(transactions)
            
            # Save transactions to database
            for txn in categorized:
                try:
                    # Get category ID
                    category_name = txn.get("category", "Miscellaneous")
                    category_id = self.categories.get(category_name, self.categories["Miscellaneous"])
                    
                    # Format date
                    date_str = txn.get("date", None)
                    try:
                        # Try to parse the date
                        date_obj = datetime.datetime.strptime(date_str, "%Y-%m-%d").date()
                    except (ValueError, TypeError):
                        # If date parsing fails, use today's date
                        date_obj = datetime.date.today()
                    
                    # Insert transaction
                    execute_query(
                        """INSERT INTO transactions 
                           (statement_id, date, description, amount, type, category_id) 
                           VALUES (%s, %s, %s, %s, %s, %s)""",
                        (
                            statement_id, 
                            date_obj,
                            txn.get("description", ""),
                            float(txn.get("amount", 0)),
                            txn.get("type", "debit").lower(),
                            category_id
                        )
                    )
                except Exception as e:
                    print(f"Error saving transaction: {e}")
            
            # Generate and save spending summary
            summary = self.generate_spending_summary(categorized)
            for category, total in summary.items():
                if total > 0:
                    try:
                        execute_query(
                            """INSERT INTO spending_summaries 
                               (statement_id, category_id, total_amount) 
                               VALUES (%s, %s, %s)""",
                            (
                                statement_id,
                                self.categories.get(category, self.categories["Miscellaneous"]),
                                total
                            )
                        )
                    except Exception as e:
                        print(f"Error saving summary: {e}")
            
            # Update statement status
            execute_query(
                """UPDATE statements 
                   SET processing_status = 'completed', processed_at = NOW() 
                   WHERE id = %s""",
                (statement_id,)
            )
            
            # Return the results
            return {
                "statement_id": statement_id,
                "transactions": categorized,
                "summary": summary
            }
            
        except Exception as e:
            # Update statement status to failed
            if 'statement_id' in locals():
                execute_query(
                    """UPDATE statements 
                       SET processing_status = 'failed' 
                       WHERE id = %s""",
                    (statement_id,)
                )
            raise e